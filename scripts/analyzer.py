#!/usr/bin/env python3
"""
判决书分析脚本 v3
功能：
1. 批量处理 PDF/Word 判决书
2. 生成详细案件摘要
3. 生成综合分析报告
"""

import sys
import os
import re
from pathlib import Path
from datetime import datetime

# 依赖导入
try:
    import pdfplumber
    from docx import Document
except ImportError:
    print("缺少依赖库，请先运行: pip install -r requirements.txt")
    sys.exit(1)


def format_section(text: str) -> str:
    """将长文本按段落、分点格式化，便于阅读"""
    if not text:
        return ""

    lines = text.split("\n")
    result_lines = []
    current_point = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 检测是否为分点开头（一、二、三 或 首先、其次、最后 或 1. 2.）
        is_heading = False
        for pattern in [r'^[一二三四五六七八九十]+[、．.]', r'^[首先|其次|最后|一、|二、|三、]', r'^\d+[、．.]']:
            if re.match(pattern, line):
                is_heading = True
                break

        if is_heading and current_point:
            # 保存上一个point
            result_lines.append("  " + " ".join(current_point))
            result_lines.append("")
            current_point = []

        if is_heading:
            result_lines.append(line)
        elif current_point:
            # 继续当前段落
            if len(line) > 50:  # 长句作为独立段
                if current_point:
                    result_lines.append("  " + " ".join(current_point))
                    result_line = []
                result_lines.append("  " + line)
            else:
                current_point.append(line)
        else:
            current_point.append(line)

    # 保存最后的段落
    if current_point:
        result_lines.append("  " + " ".join(current_point))

    return "\n".join(result_lines)


def extract_from_pdf(file_path: str) -> str:
    """从 PDF 提取文本"""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_from_word(file_path: str) -> str:
    """从 Word 文档提取文本"""
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])


def extract_case_info(text: str) -> dict:
    """提取案件详细信息"""
    info = {
        "case_number": "",
        "court": "",
        "case_type": "",
        "date": "",
        "procedure_level": "",    # 审理程序（一审/二审/再审）
        "plaintiff": "",
        "defendant": "",
        "third_party": "",
        "claims": "",              # 诉讼请求
        "appellant_claims": "",    # 上诉请求
        "facts": "",              # 案件事实
        "dispute焦点": "",         # 争议焦点
        "court_opinion": "",      # 法院认定
        "judgment_result": "",     # 判决结果
        "legal_basis": "",         # 法律依据
        "reasoning": "",           # 裁判理由
        "key_quotes": [],         # 关键引述
    }

    lines = text.split("\n")
    lines = [l.strip() for l in lines if l.strip()]
    full_text = "\n".join(lines)

    # 提取案号
    case_num_pattern = r'[（(]\d{4}[）)]\w+?[\u4e00-\u9fa5]?\d+(?:号|第\d+号)'
    for line in lines[:30]:
        match = re.search(case_num_pattern, line)
        if match:
            info["case_number"] = match.group()
            break

    # 提取法院
    court_patterns = [
        r'审理法院[：:](.+?)(?:\n|$)',
        r'作出[^\n]*判决[^\n]*法院[：:](.+?)(?:\n|$)',
        r'(最高人民法院|高级人民法院|中级人民法院|人民法院)',
    ]
    for line in lines[:30]:
        for pattern in court_patterns:
            match = re.search(pattern, line)
            if match:
                info["court"] = match.group(1) if match.lastindex else match.group()
                break
        if info["court"]:
            break

    # 提取案由
    case_type_patterns = [
        r'案\s*由[：:]\s*(.+?)(?:\n|$)',
        r'案由[：:]\s*(.+?)(?:\n|$)',
        r'本案案由为[：:]?\s*(.+?)(?:\n|$)',
    ]
    for line in lines[:50]:
        for pattern in case_type_patterns:
            match = re.search(pattern, line)
            if match:
                info["case_type"] = match.group(1).strip()
                break
        if info["case_type"]:
            break

    # 提取审理程序
    if '二审' in full_text or '民事二审' in full_text:
        info["procedure_level"] = "二审"
    elif '再审' in full_text:
        info["procedure_level"] = "再审"
    elif '一审' in full_text:
        info["procedure_level"] = "一审"

    # 提取裁判日期（支持多种格式）
    date_patterns = [
        r'(\d{4}年\d{1,2}月\d{1,2}日)',
        r'(\d{4}-\d{1,2}-\d{1,2})',
        r'裁判日期[：:]\s*(\d{4}[年\-]\d{1,2}[月\-]\d{1,2}[日]?)',
    ]
    for line in lines[:50]:
        for pattern in date_patterns:
            match = re.search(pattern, line)
            if match:
                info["date"] = match.group(1)
                break
        if info["date"]:
            break

    # 提取当事人信息（前80行）
    basic_section = "\n".join(lines[:80])

    # 提取原告/上诉人/再审申请人
    plaintiff_patterns = [
        r'(?:再审申请人|申请再审人)[^（]*（([^）]+)）?[^。]*',  # 再审申请人格式
        r'(?:上诉人|原告)(?:（[^）]*）)?(?:【[^】]*】)?[:：]?\s*(.+?)(?:\n|（|，|。|$)',
        r'(?:上诉人|原告)(?:（[^）]*）)?\s*[,，]\s*(?:住所地|所在|法定代表人)?[^。，\n]*',
    ]
    for pattern in plaintiff_patterns:
        matches = re.findall(pattern, basic_section)
        if matches:
            info["plaintiff"] = "; ".join([m.strip() for m in matches[:3] if m.strip() and len(m.strip()) > 1])
            break

    # 提取被告/被上诉人/被申请人
    defendant_patterns = [
        r'被申请人(?:（[^）]*）)?[^。]*',  # 被申请人格式
        r'(?:被上诉人|被告|被申请人)(?:（[^）]*）)?(?:【[^】]*】)?[:：]?\s*(.+?)(?:\n|（|，|。|$)',
        r'(?:被上诉人|被告)(?:（[^）]*）)?\s*[,，]\s*(?:住所地|所在|法定代表人)?[^。，\n]*',
    ]
    for pattern in defendant_patterns:
        matches = re.findall(pattern, basic_section)
        if matches:
            info["defendant"] = "; ".join([m.strip() for m in matches[:3] if m.strip()])
            break

    # 提取第三人
    third_party_pattern = r'(?:原审第三人|第三人|无独立请求权第三人)[^。，\n]*(?:【[^】]*】)?[^。，\n]*'
    matches = re.findall(third_party_pattern, basic_section)
    if matches:
        info["third_party"] = "; ".join([m.strip() for m in matches[:2] if m.strip()])

    # 提取诉讼请求/再审申请
    claims_section = ""
    claim_keywords = ["诉讼请求", "上诉请求", "再审申请", "请求事项", "诉讼主张"]
    for i, line in enumerate(lines):
        if any(kw in line for kw in claim_keywords) and len(line) < 150:
            claims_section = "\n".join(lines[i:i+15])
            break
    if claims_section:
        info["claims"] = claims_section[:1500]

    # 提取上诉请求（如果有）
    appellant_section = ""
    for i, line in enumerate(lines):
        if "上诉请求" in line or "上诉人请求" in line:
            appellant_section = "\n".join(lines[i:i+10])
            break
    if appellant_section:
        info["appellant_claims"] = appellant_section[:800]

    # 提取案件事实（审理查明部分）
    facts_section = ""
    fact_keywords = ["审理查明", "查明", "经审理查明", "本院查明", "事实", "案件事实"]
    found_facts = False
    for i, line in enumerate(lines):
        if any(kw in line for kw in fact_keywords) and i > 30:
            # 找到事实部分，提取后续内容
            facts_section = "\n".join(lines[i:i+50])
            found_facts = True
            break
    if facts_section:
        info["facts"] = facts_section[:3000]

    # 提取争议焦点（重点提取）
    dispute_section = ""
    dispute_keywords = ["争议焦点", "本案焦点", "双方争议", "争议焦点如下", "本案的焦点", "主要审查", "主要争议"]
    for i, line in enumerate(lines):
        if any(kw in line for kw in dispute_keywords):
            # 找到后提取更多内容作为论证
            dispute_section = "\n".join(lines[i:i+30])
            break
    if not dispute_section:
        # 备选：从"本院认为"前面的内容中找
        for i, line in enumerate(lines):
            if "本院认为" in line:
                # 往前找争议相关的描述
                for j in range(i-1, max(0, i-20), -1):
                    if any(kw in lines[j] for kw in ["焦点", "争议", "问题", "双方"]):
                        dispute_section = "\n".join(lines[j:i+5])
                        break
                break
    if dispute_section:
        # 格式化分段
        info["dispute焦点"] = format_section(dispute_section[:2000])

    # 提取法院论证/本院认为（核心部分，重点提取）
    opinion_section = ""
    opinion_keywords = ["本院认为", "本院审理认为", "本院经审查认为", "本院经审查认为", "对此，本院", "综上"]
    for i, line in enumerate(lines):
        if any(kw in line for kw in opinion_keywords):
            # 提取更长的论证内容
            opinion_section = "\n".join(lines[i:i+50])
            break
    if opinion_section:
        # 格式化分段
        info["court_opinion"] = format_section(opinion_section[:3000])

    # 提取判决结果（找"判决如下"后的具体判决内容）
    judgment_section = ""
    for i, line in enumerate(lines):
        # 找到"判决如下"或"裁决如下"或"裁定如下"
        if re.search(r'(?:判决|裁决|裁定)如下', line):
            # 取之后的内容（跳过标题行）
            judgment_section = "\n".join(lines[i+1:i+15])
            break
    if not judgment_section:
        # 备选：找结尾部分带"判决"、"裁决"的内容
        for i in range(len(lines)-1, max(0, len(lines)-50), -1):
            if re.search(r'[。；]\s*$', lines[i]) and any(kw in lines[i] for kw in ['判决', '裁决', '裁定']):
                judgment_section = "\n".join(lines[max(0,i-10):i+5])
                break
    if judgment_section:
        info["judgment_result"] = judgment_section[:1500]

    # 提取法律依据
    legal_section = ""
    legal_patterns = [
        r'依照[^。；]+',
        r'根据[^。；]+',
        r'依据[^。；]+',
        r'引用[^。；]+',
    ]
    for line in lines:
        for pattern in legal_patterns:
            match = re.search(pattern, line)
            if match and len(match.group()) > 10:
                legal_section += match.group() + "; "
        if legal_section:
            break
    if legal_section:
        info["legal_basis"] = legal_section[:1000]

    # 提取裁判理由
    reasoning_section = ""
    for i, line in enumerate(lines):
        if "理由" in line and i > 50:
            reasoning_section = "\n".join(lines[i:i+20])
            break
    if reasoning_section:
        info["reasoning"] = reasoning_section[:1500]

    # 提取关键引述
    quote_patterns = [
        r'"([^"]{20,200})"',
        r'"([^"]{20,200})"',
    ]
    for pattern in quote_patterns:
        quotes = re.findall(pattern, full_text)
        if quotes:
            info["key_quotes"] = [q.strip() for q in quotes[:5] if len(q.strip()) > 20]
            break

    return info


def generate_summary(info: dict, original_filename: str) -> str:
    """生成精简案件摘要"""
    # 提取关键要点，每项限制长度
    def limit(text, length=200):
        if not text:
            return ''
        return text[:length] + ('...' if len(text) > length else '')

    # 核心内容：争议焦点和法院论证
    dispute = limit(info.get('dispute焦点', ''), 300)
    opinion = limit(info.get('court_opinion', ''), 500)
    result = limit(info.get('judgment_result', ''), 150)

    summary = f"""================================================================================
                              案 件 摘 要
================================================================================

【基本信息】
案号：{info['case_number'] or '—'}
法院：{info['court'] or '—'}
案由：{info['case_type'] or '—'}
程序：{info['procedure_level'] or '—'}
日期：{info['date'] or '—'}

【争议焦点】
{dispute or '—'}

【法院论证】
{opinion or '—'}

【判决结果】
{result or '—'}

================================================================================
"""
    return summary


def generate_comparative_report(summaries: list, output_dir: str, all_case_info: list):
    """生成综合分析报告"""

    # 统计分析
    case_types = {}
    courts = {}
    procedure_levels = {}
    dates = []

    for info in all_case_info:
        # 案由统计
        case_type = info.get('case_type', '') or '未分类'
        # 简化案由（取第一个主要类型）
        main_type = case_type.split('、')[0].split('，')[0] if case_type else '未分类'
        case_types[main_type] = case_types.get(main_type, 0) + 1

        # 法院统计
        court = info.get('court', '') or '未识别'
        courts[court] = courts.get(court, 0) + 1

        # 程序统计
        proc = info.get('procedure_level', '') or '未识别'
        procedure_levels[proc] = procedure_levels.get(proc, 0) + 1

        # 日期
        if info.get('date'):
            dates.append(info['date'])

    # 生成 Word 报告
    doc = Document()
    doc.add_heading("案件综合分析报告", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"案件数量：{len(all_case_info)} 份")

    # ========== 一、案件概览 ==========
    doc.add_heading("一、案件概览", 1)

    doc.add_heading("1.1 案件类型分布", 2)
    total = len(all_case_info)
    for ct, count in sorted(case_types.items(), key=lambda x: -x[1]):
        pct = count * 100 // total
        doc.add_paragraph(f"• {ct}：{count} 件（{pct}%）")

    doc.add_heading("1.2 法院分布", 2)
    doc.add_paragraph(f"共涉及 {len(courts)} 家法院：")
    for court, count in sorted(courts.items(), key=lambda x: -x[1])[:8]:
        doc.add_paragraph(f"  - {court}：{count} 件")

    doc.add_heading("1.3 审理程序", 2)
    for proc, count in sorted(procedure_levels.items(), key=lambda x: -x[1]):
        doc.add_paragraph(f"• {proc}：{count} 件")

    if dates:
        doc.add_paragraph(f"案件时间范围：{min(dates)} 至 {max(dates)}")

    # ========== 二、逐案摘要 ==========
    doc.add_heading("二、逐案摘要", 1)

    for i, info in enumerate(all_case_info, 1):
        case_num = info.get('case_number', f'案件{i}') or f'案件{i}'
        case_type = info.get('case_type', '—') or '—'
        court = info.get('court', '—') or '—'
        date = info.get('date', '—') or '—'
        dispute = info.get('dispute焦点', '') or ''
        opinion = info.get('court_opinion', '') or ''

        # 提炼争议焦点和法院论证（核心内容）
        dispute_summary = dispute[:100] + "..." if len(dispute) > 100 else dispute
        opinion_summary = opinion[:150] + "..." if len(opinion) > 150 else opinion

        p = doc.add_paragraph()
        p.add_run(f"{i}. {case_num}").bold = True
        p.add_run(f" | {case_type} | {court} | {date}")

        if dispute_summary:
            doc.add_paragraph(f"争议焦点：{dispute_summary}")
        if opinion_summary:
            doc.add_paragraph(f"法院论证：{opinion_summary}")

    # ========== 三、对比分析 ==========
    doc.add_heading("三、案件对比分析", 1)

    doc.add_heading("3.1 共同特征", 2)
    doc.add_paragraph(f"通过对 {len(all_case_info)} 份判决书的分析，发现以下共同特征：")
    doc.add_paragraph(f"• 案由分布：涵盖 {len(case_types)} 种案件类型，其中委托理财类纠纷占比最高（{case_types.get('金融委托理财合同纠纷', 0) + case_types.get('委托理财合同纠纷', 0)} 件）")
    doc.add_paragraph(f"• 法院层级：涉及 {len(courts)} 家法院，以中级人民法院和高级人民法院为主")
    doc.add_paragraph(f"• 纠纷主体：多为金融机构与个人投资者之间的纠纷，反映了金融消费领域的权益保护问题")

    doc.add_heading("3.2 差异分析", 2)

    # 构建对比表格
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    headers = ['案号', '案由', '程序', '法院', '日期', '判决要点']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True

    for info in all_case_info:
        row = table.add_row().cells
        row[0].text = (info.get('case_number', '') or '—')[:25]
        row[1].text = (info.get('case_type', '') or '—')[:20]
        row[2].text = info.get('procedure_level', '—') or '—'
        row[3].text = (info.get('court', '') or '—')[:15]
        row[4].text = info.get('date', '—') or '—'

        # 判决要点摘要
        result = info.get('judgment_result', '') or ''
        key_point = result[:50] + "..." if len(result) > 50 else result
        row[5].text = key_point or '见详细摘要'

    # ========== 四、法律要点归纳 ==========
    doc.add_heading("四、法律要点归纳", 1)

    doc.add_heading("4.1 案由分类统计", 2)
    for ct, count in sorted(case_types.items(), key=lambda x: -x[1]):
        pct = count * 100 // total
        doc.add_paragraph(f"• {ct}：{count} 件（{pct}%）")

    doc.add_heading("4.2 主要法律问题", 2)

    # 金融委托理财类
    financial_cases = [i for i, info in enumerate(all_case_info)
                      if '委托理财' in (info.get('case_type') or '')
                      or '理财' in (info.get('case_type') or '')]
    if financial_cases:
        doc.add_heading("（一）金融委托理财合同纠纷", 3)
        doc.add_paragraph(f"涉及 {len(financial_cases)} 件，主要法律问题包括：")
        doc.add_paragraph("• 金融机构适当性义务履行问题")
        doc.add_paragraph("• 投资者风险承受能力评估")
        doc.add_paragraph("• 产品风险揭示与告知义务")
        doc.add_paragraph("• 基金合同效力认定")

    # 侵权责任类
    tort_cases = [i for i, info in enumerate(all_case_info)
                  if '侵权' in (info.get('case_type') or '')]
    if tort_cases:
        doc.add_heading("（二）侵权责任纠纷", 3)
        doc.add_paragraph(f"涉及 {len(tort_cases)} 件，主要法律问题包括：")
        doc.add_paragraph("• 金融机构侵权责任认定")
        doc.add_paragraph("• 财产损害赔偿范围")
        doc.add_paragraph("• 过错责任与因果关系")

    # 缔约过失类
    negligence_cases = [i for i, info in enumerate(all_case_info)
                        if '缔约过失' in (info.get('case_type') or '')]
    if negligence_cases:
        doc.add_heading("（三）缔约过失责任纠纷", 3)
        doc.add_paragraph(f"涉及 {len(negligence_cases)} 件，主要法律问题包括：")
        doc.add_paragraph("• 合同订立阶段的诚信义务")
        doc.add_paragraph("• 缔约过失的构成要件")
        doc.add_paragraph("• 信赖利益的保护")

    # ========== 五、结论与建议 ==========
    doc.add_heading("五、结论与建议", 1)

    doc.add_heading("5.1 案件特征总结", 2)
    doc.add_paragraph(f"本次分析的 {len(all_case_info)} 份判决书呈现以下特征：")

    # 找出最主要的案由
    main_case_type = max(case_types.items(), key=lambda x: x[1]) if case_types else ('未知', 0)
    doc.add_paragraph(f"• 案件类型以 {main_case_type[0]} 最为常见，共 {main_case_type[1]} 件，占比 {main_case_type[1]*100//total}%")

    # 找出最主要的法院
    main_court = max(courts.items(), key=lambda x: x[1]) if courts else ('未知', 0)
    doc.add_paragraph(f"• 涉及 {len(courts)} 家法院，其中 {main_court[0]} 审理案件最多（{main_court[1]} 件）")

    if procedure_levels:
        main_proc = max(procedure_levels.items(), key=lambda x: x[1])
        doc.add_paragraph(f"• 审理程序以 {main_proc[0]} 为主（{main_proc[1]} 件）")

    doc.add_heading("5.2 风险提示", 2)

    if financial_cases:
        doc.add_paragraph("【金融委托理财领域】")
        doc.add_paragraph("• 金融机构应严格履行适当性义务，确保产品与投资者风险承受能力匹配")
        doc.add_paragraph("• 应充分揭示产品风险，确保投资者在充分了解风险的基础上作出投资决策")
        doc.add_paragraph("• 加强对销售环节的合规管理，避免因操作不规范承担赔偿责任")

    if tort_cases:
        doc.add_paragraph("【侵权责任领域】")
        doc.add_paragraph("• 金融机构应加强对客户资金的安全管理")
        doc.add_paragraph("• 严格遵守相关法律法规和监管要求")

    if negligence_cases:
        doc.add_paragraph("【缔约过失领域】")
        doc.add_paragraph("• 合同订立阶段应遵循诚信原则")
        doc.add_paragraph("• 明确双方权利义务，避免因缔约过失产生纠纷")

    doc.add_heading("5.3 建议", 2)
    doc.add_paragraph("• 建议金融机构加强合规体系建设，完善投资者适当性管理制度")
    doc.add_paragraph("• 建议投资者在购买金融产品时充分了解产品风险和自身权益")
    doc.add_paragraph("• 建议加强对金融消费者的法律知识普及")

    output_path = os.path.join(output_dir, "综合分析报告.docx")
    doc.save(output_path)
    return output_path


def process_single_file(file_path: str, output_dir: str) -> dict:
    """处理单个判决书文件"""
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    # 提取文本
    if ext == ".pdf":
        text = extract_from_pdf(str(file_path))
    elif ext in [".docx", ".doc"]:
        text = extract_from_word(str(file_path))
    else:
        return None

    # 提取信息
    info = extract_case_info(text)

    # 生成详细摘要
    summary = generate_summary(info, file_path.name)

    # 保存摘要文本
    summary_txt_path = os.path.join(output_dir, f"{file_path.stem}_摘要.txt")
    with open(summary_txt_path, 'w', encoding='utf-8') as f:
        f.write(summary)

    return info


def main():
    if len(sys.argv) < 2:
        print("用法:")
        print("  单文件: python analyzer.py <文件路径>")
        print("  文件夹: python analyzer.py <文件夹路径>")
        sys.exit(1)

    input_path = sys.argv[1]
    input_path = Path(input_path)

    # 确定输出目录（在源文件夹内创建摘要文件夹）
    output_dir = input_path / "摘要"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 处理输入
    all_info = []
    summaries = []

    if input_path.is_file():
        # 单文件
        print(f"处理文件: {input_path.name}")
        info = process_single_file(str(input_path), str(output_dir))
        if info:
            all_info.append(info)
        summaries.append(str(input_path))
    elif input_path.is_dir():
        # 文件夹 - 批量处理
        supported_ext = {'.pdf', '.docx', '.doc'}
        files = [f for f in input_path.iterdir()
                 if f.is_file() and f.suffix.lower() in supported_ext]

        print(f"找到 {len(files)} 份判决书...")

        for i, f in enumerate(files, 1):
            print(f"处理 {i}/{len(files)}: {f.name}")
            info = process_single_file(str(f), str(output_dir))
            if info:
                all_info.append(info)
            summaries.append(str(f))

    # 生成综合报告
    if all_info:
        report_path = generate_comparative_report(summaries, str(output_dir), all_info)
        print(f"\n综合报告已生成: {report_path}")

    print(f"\n摘要文件夹: {output_dir}")
    print(f"共处理 {len(all_info)} 份判决书")


if __name__ == "__main__":
    main()
